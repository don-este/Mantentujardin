import streamlit as st
from docx import Document
from io import BytesIO
import datetime

# --- CONFIGURACIÓN DE ACCESO ---
def check_password():
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False
    if not st.session_state["password_correct"]:
        st.title("🔐 Acceso Mantentujardin")
        user = st.text_input("Usuario")
        password = st.text_input("Contraseña", type="password")
        if st.button("Ingresar"):
            if user == "esteban" and password == "jardin2024": # <--- Cambia esto
                st.session_state["password_correct"] = True
                st.rerun()
            else:
                st.error("Error de credenciales")
        return False
    return True

if check_password():
    st.title("🌿 Registro de Trabajos")
    
    # Formulario profesional
    with st.form("main_form"):
        nombre = st.text_input("Nombre del Cliente")
        direccion = st.text_input("Dirección")
        servicio = st.multiselect("Servicios realizados", 
                                 ["Mantención Piscina", "Corte de Césped", "Poda", "Otros"])
        modalidad = st.radio("Tipo de cobro", ["Mensual", "Por Visita"])
        valor = st.number_input("Valor total ($)", min_value=0)
        
        if st.form_submit_button("Registrar y Crear Boleta"):
            if nombre and valor > 0:
                # Crear Word
                doc = Document()
                doc.add_heading('BOLETA DE SERVICIO', 0)
                doc.add_paragraph(f"Fecha: {datetime.date.today()}")
                doc.add_paragraph(f"Cliente: {nombre}\nDirección: {direccion}")
                doc.add_paragraph(f"Servicios: {', '.join(servicio)}")
                doc.add_paragraph(f"Modalidad: {modalidad}")
                doc.add_heading(f"TOTAL: ${valor:,}", level=1)
                
                # Preparar descarga
                buf = BytesIO()
                doc.save(buf)
                st.session_state.documento = buf.getvalue()
                st.success("✅ ¡Boleta lista!")

    if "documento" in st.session_state:
        st.download_button("📥 Descargar Boleta Word", st.session_state.documento, "Boleta.docx")